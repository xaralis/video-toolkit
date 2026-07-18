#!/usr/bin/env python3
"""Render `projects/<name>/SCREENPLAY.md` → `NATACENI.docx` — printable field
shooting cue-cards the speaker reads straight off the page while filming.

The document has two parts:

  * an **overview page** (portrait A4): the b-roll shot list, general speaking
    tips, and a numbered outline of the takes in shooting order;
  * one **giant-font landscape page per spoken take**: the exact words to say,
    sized big enough to read from a phone or a printout at arm's length.

A "take" is one continuous spoken passage: a face-to-camera clip plus every
following b-roll segment whose audio L-cuts from it (the voice keeps running
while the picture cuts to b-roll). The speaker records the whole take in one go.

This is the final step of `/toolkit:narrate` — run it after SCREENPLAY.md is
authored so the field cards always match the current screenplay.

Usage:
    python3 -m video_toolkit.render_shooting_cards <project-name>
    python3 -m video_toolkit.render_shooting_cards --all   # every project with a SCREENPLAY.md
"""
from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

from video_toolkit.paths import workspace_root

# General speaking tips — boilerplate, same for every project. Kept in Czech
# because the cards are read in the field by Czech-speaking collaborators.
SPEAKING_TIPS = [
    "Mluv přirozeně, jako bys to vysvětloval kamarádovi – ne jako když čteš z papíru.",
    "Když má promluva víc vět za sebou, řekni je jedním tahem, bez pauz uprostřed.",
    "Klidně to zkus vícekrát za sebou – použijeme nejlepší verzi.",
    "Dívej se do kamery, ne do telefonu ani do papíru.",
    "Každá promluva je na samostatné stránce dál v tomto dokumentu – natáčej je popořadě.",
]

_SEG_RE = re.compile(r"^##\s+seg-\d+\s+\[(?P<bracket>[^\]]+)\]\s*$", re.MULTILINE)
_SPOKEN_NEW_RE = re.compile(r"\*\*Spoken intent:\*\*\s*(?P<text>.+)")
_SPOKEN_LCUT_RE = re.compile(r"\*\*Spoken intent \([^)]*\):\*\*\s*(?P<text>.+)")
_VISUAL_RE = re.compile(r"\*\*Visual intent:\*\*\s*(?P<text>.+)")


@dataclass
class Take:
    label: str
    lines: list[str] = field(default_factory=list)


@dataclass
class Screenplay:
    title: str
    takes: list[Take] = field(default_factory=list)
    broll: list[str] = field(default_factory=list)


def _split_frontmatter(text: str) -> tuple[dict[str, str], str]:
    """Return (frontmatter dict, body) — frontmatter is the leading --- block."""
    fm: dict[str, str] = {}
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            block = text[3:end]
            for line in block.splitlines():
                if ":" in line:
                    k, _, v = line.partition(":")
                    fm[k.strip()] = v.strip().strip('"')
            return fm, text[end + 4:]
    return fm, text


def _title(fm: dict[str, str], body: str) -> str:
    """Prefer the '# Screenplay — X' / '# Natáčení — X' heading, uppercased."""
    m = re.search(r"^#\s+(?:Screenplay|Natáčení)\s+[—-]\s+(?P<t>.+)$", body, re.MULTILINE)
    if m:
        return m.group("t").strip().upper()
    if fm.get("chevron"):
        return fm["chevron"].strip().upper()
    return "NATÁČENÍ"


def parse_screenplay(md: str) -> Screenplay:
    fm, body = _split_frontmatter(md)
    sp = Screenplay(title=_title(fm, body))

    matches = list(_SEG_RE.finditer(body))
    for i, m in enumerate(matches):
        bracket = m.group("bracket")
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(body)
        block = body[start:end]

        kind = bracket.split("·")[0].strip()
        parts = [p.strip() for p in bracket.split("·")]
        label = parts[-1] if len(parts) > 1 else kind

        spoken_new = _SPOKEN_NEW_RE.search(block)
        spoken_lcut = _SPOKEN_LCUT_RE.search(block)
        visual = _VISUAL_RE.search(block)

        if spoken_new:
            # A new face-to-camera take begins here.
            sp.takes.append(Take(label=label, lines=[spoken_new.group("text").strip()]))
        elif spoken_lcut and sp.takes:
            # Audio L-cuts from the current take — same continuous passage.
            sp.takes[-1].lines.append(spoken_lcut.group("text").strip())

        if kind == "broll" and visual:
            sp.broll.append(visual.group("text").strip())

    return sp


import math

# --- One-page-per-take fit (a take is a cue card; the speaker must never turn a
# page mid-speech, so each take is GUARANTEED to fit its landscape A4 page — and
# fills it as fully as possible: the font is the LARGEST size that still fits). --
#
# The size is derived from the page geometry, not a char-count guess: estimate
# the take's rendered height at each candidate size and keep the largest that
# fits. The search is per-point (not a coarse ladder) so a take fills its page
# instead of dropping a whole tier and leaving whitespace. The model stays a
# touch pessimistic (average glyph a hair wider than common bold sans, line
# spacing above the 1.15 actually applied, 5 % height slack) so the guarantee
# survives Word's own line-breaking and font substitution. Geometry MUST match
# the per-take section in build_docx (landscape A4, 20 mm sides, 16 mm top/bottom).
_CARD_USABLE_W_PT = (297 - 2 * 20) / 25.4 * 72  # 257 mm ≈ 728.5 pt
_CARD_USABLE_H_PT = (210 - 2 * 16) / 25.4 * 72  # 178 mm ≈ 504.6 pt
_TAG_BLOCK_PT = 14 * 1.20 + 16              # "PROMLUVA n — label" line + its space_after
_AVG_CHAR_W = 0.56          # bold-sans advance as a fraction of em (a hair over Arial/Calibri bold)
_MODEL_LINE_SPACING = 1.20  # above the 1.15 actually applied → built-in slack
_PARA_AFTER = 0.30          # space after each fragment, as a fraction of the font size
_FIT_SAFETY = 0.95          # never fill more than 95 % of the usable height (buffer for model error)
_MAX_TAKE_PT = 80           # never larger than this even for a one-word take
_MIN_TAKE_PT = 22           # readable floor; below this we warn instead of shrinking further


def _estimate_take_height_pt(lines: list[str], size: int) -> float:
    """Conservative rendered height (pt) of a take at the given font size."""
    chars_per_line = max(1, int(_CARD_USABLE_W_PT / (_AVG_CHAR_W * size)))
    height = _TAG_BLOCK_PT
    for line in lines:
        wrapped = max(1, math.ceil(len(line) / chars_per_line))
        height += wrapped * size * _MODEL_LINE_SPACING + size * _PARA_AFTER
    return height


def _fit_take_font(lines: list[str]) -> tuple[int, bool]:
    """Largest per-point size whose take fits one landscape page.

    Returns ``(size_pt, fits)``. ``fits`` is False only when even the minimum
    size overflows — the take is too long for a single card and the caller warns
    the author to split it."""
    budget = _CARD_USABLE_H_PT * _FIT_SAFETY
    for size in range(_MAX_TAKE_PT, _MIN_TAKE_PT - 1, -1):
        if _estimate_take_height_pt(lines, size) <= budget:
            return size, True
    return _MIN_TAKE_PT, False


def build_docx(sp: Screenplay, out_path: Path) -> list[int]:
    """Write the cue-card document; return the 1-based indices of any takes too
    long to fit a single page even at the minimum font size (empty when all fit)."""
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    A4_W, A4_H = Mm(210), Mm(297)
    doc = Document()

    # --- Overview page (portrait A4) ---
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width, sec.page_height = A4_W, A4_H
    sec.top_margin = sec.bottom_margin = Mm(16)
    sec.left_margin = sec.right_margin = Mm(18)

    def heading(text: str, size: int, space_before: int = 10) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(size)

    title_p = doc.add_paragraph()
    tr = title_p.add_run(sp.title)
    tr.bold = True
    tr.font.size = Pt(26)
    sub = doc.add_paragraph()
    sr = sub.add_run("Pokyny k natočení — vytiskni a vezmi s sebou na natáčení.")
    sr.font.size = Pt(14)

    if sp.broll:
        heading("PŘESTŘIHY (BROLL) — natoč bez mluvení", 16)
        for shot in sp.broll:
            b = doc.add_paragraph(style="List Bullet")
            b.add_run(shot).font.size = Pt(13)

    heading("OBECNÉ POKYNY PRO PROMLUVY", 16)
    for tip in SPEAKING_TIPS:
        b = doc.add_paragraph(style="List Bullet")
        b.add_run(tip).font.size = Pt(13)

    if sp.takes:
        heading("PŘEHLED PROMLUV (v tomto pořadí je natoč)", 16)
        for i, take in enumerate(sp.takes, 1):
            first = take.lines[0] if take.lines else ""
            preview = first if len(first) <= 55 else first[:54].rstrip() + "…"
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"{i}. {preview}  ({take.label})")
            r.bold = True
            r.font.size = Pt(12)

    # --- One giant-font landscape page per take (guaranteed to fit) ---
    overflow: list[int] = []
    for i, take in enumerate(sp.takes, 1):
        s = doc.add_section(WD_SECTION.NEW_PAGE)
        s.orientation = WD_ORIENT.LANDSCAPE
        s.page_width, s.page_height = A4_H, A4_W  # swapped for landscape
        s.top_margin = s.bottom_margin = Mm(16)
        s.left_margin = s.right_margin = Mm(20)

        tag = doc.add_paragraph()
        tr = tag.add_run(f"PROMLUVA {i} — {take.label}")
        tr.bold = True
        tr.font.name = "Arial"
        tr.font.size = Pt(14)
        tag.paragraph_format.space_before = Pt(0)
        tag.paragraph_format.space_after = Pt(16)
        tag.paragraph_format.line_spacing = 1.15

        size, fits = _fit_take_font(take.lines)
        if not fits:
            overflow.append(i)
        for line in take.lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(size * _PARA_AFTER)
            pf.line_spacing = 1.15
            r = p.add_run(line)
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(size)

    doc.save(str(out_path))
    return overflow


def project_dir(name: str) -> Path:
    p = workspace_root() / "projects" / name
    if not p.is_dir():
        sys.exit(f"error: project directory not found: {p}")
    return p


def render_project(name: str) -> Path:
    p = project_dir(name)
    md = p / "SCREENPLAY.md"
    if not md.is_file():
        sys.exit(f"error: {md} not found")
    sp = parse_screenplay(md.read_text(encoding="utf-8"))
    if not sp.takes:
        sys.exit(f"error: no spoken takes found in {md} — nothing to put on cards")
    out = p / "NATACENI.docx"
    overflow = build_docx(sp, out)
    for i in overflow:
        label = sp.takes[i - 1].label
        print(
            f"⚠ warning: take {i} ({label!r}) is too long to fit one page even at "
            f"{_MIN_TAKE_PT}pt — split it into shorter takes so the speaker never "
            f"turns a page mid-speech.",
            file=sys.stderr,
        )
    return out


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Render SCREENPLAY.md → NATACENI.docx (printable field shooting cards)."
    )
    g = parser.add_mutually_exclusive_group(required=True)
    g.add_argument("project", nargs="?", help="Project directory under projects/")
    g.add_argument("--all", action="store_true", help="Render every project with a SCREENPLAY.md")
    args = parser.parse_args()

    try:
        import docx  # noqa: F401
    except ModuleNotFoundError:
        sys.exit("error: python-docx not installed — run `pip install -e toolkit/` (or `-e .` in core)")

    if args.all:
        projects_dir = workspace_root() / "projects"
        targets = sorted(
            p.name for p in projects_dir.iterdir()
            if p.is_dir() and (p / "SCREENPLAY.md").is_file()
        )
        if not targets:
            sys.exit("error: no projects with SCREENPLAY.md found")
        for name in targets:
            print(f"✔ {render_project(name)}")
    else:
        print(f"✔ {render_project(args.project)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
